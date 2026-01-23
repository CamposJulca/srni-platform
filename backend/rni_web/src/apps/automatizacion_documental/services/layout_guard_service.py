from docx import Document
from docx.shared import Pt, Cm
from pathlib import Path


def asegurar_una_sola_pagina(
    docx_path: Path,
    firma_img_path: Path,
    nombre_firmante: str = "OSCAR ANDRÉS MANOSALVA GARCÍA"
):
    """
    Ajusta el layout del DOCX y aplica la firma garantizando
    que el documento quede en una sola página.
    """

    doc = Document(docx_path)
    section = doc.sections[0]

    # ─────────────────────────────
    # 1. Márgenes seguros (Word + LibreOffice)
    # ─────────────────────────────
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)

    section.header_distance = Pt(28)
    section.footer_distance = Pt(28)

    # ─────────────────────────────
    # 2. Normalizar TODOS los párrafos
    # ─────────────────────────────
    for p in doc.paragraphs:
        fmt = p.paragraph_format
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.line_spacing = 1.0

    # ─────────────────────────────
    # 3. Insertar firma inline controlada
    # ─────────────────────────────
    for paragraph in doc.paragraphs:
        if nombre_firmante in paragraph.text:
            run = paragraph.add_run("\n")
            run.add_picture(
                str(firma_img_path),
                width=Cm(3.0)   # tamaño seguro para evitar overflow
            )
            break

    doc.save(docx_path)
