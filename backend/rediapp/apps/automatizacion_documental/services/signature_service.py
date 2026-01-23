# # # # from docx import Document
# # # # from docx.shared import Cm
# # # # from pathlib import Path


# # # # def firmar_documento(docx_path: Path, firma_img_path: Path):
# # # #     """
# # # #     Inserta una firma gráfica antes del nombre del firmante
# # # #     en un documento Word (.docx).
# # # #     """

# # # #     doc = Document(docx_path)

# # # #     for i, paragraph in enumerate(doc.paragraphs):
# # # #         if "OSCAR ANDRÉS MANOSALVA GARCÍA" in paragraph.text:
# # # #             firma_paragraph = doc.paragraphs[i].insert_paragraph_before()
# # # #             run = firma_paragraph.add_run()
# # # #             run.add_picture(str(firma_img_path), width=Cm(4))
# # # #             break

# # # #     doc.save(docx_path)


# # # from docx import Document
# # # from docx.shared import Cm
# # # from pathlib import Path


# # # def firmar_documento(docx_path: Path, firma_img_path: Path):
# # #     doc = Document(docx_path)

# # #     for i, paragraph in enumerate(doc.paragraphs):
# # #         if "OSCAR ANDRÉS MANOSALVA GARCÍA" in paragraph.text:
# # #             firma_paragraph = doc.paragraphs[i].insert_paragraph_before()
# # #             run = firma_paragraph.add_run()
# # #             run.add_picture(str(firma_img_path), width=Cm(4))
# # #             break

# # #     doc.save(docx_path)


# # from docx import Document
# # from docx.shared import Cm
# # from pathlib import Path


# # def firmar_documento(docx_path: Path, firma_img_path: Path):
# #     doc = Document(docx_path)

# #     for paragraph in doc.paragraphs:
# #         if "OSCAR ANDRÉS MANOSALVA GARCÍA" in paragraph.text:
# #             run = paragraph.add_run("\n")
# #             run.add_picture(str(firma_img_path), width=Cm(3.5))
# #             break

# #     doc.save(docx_path)

# from docx import Document
# from docx.shared import Cm
# from pathlib import Path


# def firmar_documento(docx_path: Path, firma_img_path: Path):
#     """
#     Inserta la firma en el pie de página.
#     Esto garantiza que el documento NO cambie de página.
#     """

#     doc = Document(docx_path)
#     section = doc.sections[0]
#     footer = section.footer

#     # Limpiar párrafos vacíos del footer
#     for p in footer.paragraphs:
#         if not p.text.strip():
#             footer._element.remove(p._element)

#     p = footer.add_paragraph()
#     p.alignment = 1  # centrado

#     run = p.add_run()
#     run.add_picture(str(firma_img_path), width=Cm(3.5))

#     doc.save(docx_path)

from docx import Document
from docx.shared import Cm
from pathlib import Path


from docx import Document
from docx.shared import Cm
from pathlib import Path


def firmar_documento(docx_path: Path, firma_img_path: Path):
    doc = Document(docx_path)

    for p in doc.paragraphs:
        if "OSCAR ANDRÉS MANOSALVA GARCÍA" in p.text:
            run = p.add_run("\n")
            run.add_picture(str(firma_img_path), width=Cm(4))
            break

    doc.save(docx_path)
